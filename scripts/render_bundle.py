#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
TEMPLATE_DIR = SKILL_DIR / "assets" / "templates"
APPLICATION_TEMPLATE = TEMPLATE_DIR / "template_application.docx"
MANUAL_TEMPLATE = TEMPLATE_DIR / "template_manual.docx"
CODE_TEMPLATE = TEMPLATE_DIR / "template_code.docx"

if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from render_mockups import render_mockups  # noqa: E402
from template_rewriter import build_code_lines, build_manual_sequence, build_spec  # noqa: E402


RESIDUE_KEYWORDS = ["校园小动物", "campuspet", "AnimalService", "CampusPet", "动物信息录入"]


def sanitize_name(name: str) -> str:
    return "".join("_" if c in '<>:"/\\|?*' else c for c in name).strip() or "未命名系统"


def set_header_text(document: Document, text: str) -> None:
    for section in document.sections:
        if not section.header.paragraphs:
            section.header.add_paragraph()
        section.header.paragraphs[0].text = text


def fill_application_tables(document: Document, spec: dict, source_line_count: int) -> None:
    defaults = spec["defaults"]
    t1, t2, t3 = document.tables
    values1 = [defaults["rights_acquisition"], spec["software_name"], spec["version"], defaults["rights_scope"]]
    for row_index, value in enumerate(values1, start=1):
        t1.rows[row_index].cells[1].text = str(value)
    values2 = [
        defaults["software_type"],
        defaults["software_nature"],
        defaults["development_mode"],
        spec["development_date"],
        defaults["publication_status"],
    ]
    for row_index, value in enumerate(values2, start=1):
        t2.rows[row_index].cells[1].text = str(value)
    values3 = [
        defaults["dev_hardware"],
        defaults["run_hardware"],
        defaults["dev_os"],
        defaults["dev_tools"],
        defaults["run_platform"],
        defaults["support_software"],
        defaults["languages"],
        str(source_line_count),
        spec["purpose"],
        f"面向{spec['industry']}领域，主要服务于{spec['target_users']}。",
        spec["main_features"],
        spec["technical_highlights"],
    ]
    for row_index, value in enumerate(values3, start=1):
        t3.rows[row_index].cells[1].text = str(value)


def rewrite_manual_paragraphs(document: Document, manual_texts: list[str]) -> None:
    non_empty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    if len(manual_texts) > len(non_empty):
        raise ValueError("Generated manual text exceeds template paragraph count.")
    for index, paragraph in enumerate(non_empty):
        paragraph.text = manual_texts[index] if index < len(manual_texts) else ""


def replace_docx_media(docx_path: Path, image_dir: Path) -> None:
    fd, temp_name = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    temp_file = Path(temp_name)
    replacements = {f"word/media/{path.name}": path for path in image_dir.iterdir() if path.is_file()}
    with zipfile.ZipFile(docx_path, "r") as source_zip, zipfile.ZipFile(temp_file, "w") as target_zip:
        for item in source_zip.infolist():
            if item.filename in replacements:
                target_zip.writestr(item, replacements[item.filename].read_bytes())
            else:
                target_zip.writestr(item, source_zip.read(item.filename))
    shutil.move(str(temp_file), str(docx_path))


def clear_document_body(document: Document) -> None:
    body = document._element.body
    section_node = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            section_node = child
    for child in list(body):
        if child is not section_node:
            body.remove(child)


def write_code_document(document: Document, code_lines: list[str]) -> None:
    clear_document_body(document)
    style_name = None
    try:
        document.styles["No Spacing"]
        style_name = "No Spacing"
    except KeyError:
        style_name = None
    for line in code_lines:
        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def scan_for_residue(docx_path: Path) -> list[str]:
    hits: list[str] = []
    with zipfile.ZipFile(docx_path) as archive:
        text_parts = []
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                try:
                    text_parts.append(archive.read(name).decode("utf-8", errors="ignore"))
                except Exception:
                    continue
    merged = "\n".join(text_parts).lower()
    for keyword in RESIDUE_KEYWORDS:
        if keyword.lower() in merged:
            hits.append(keyword)
    return hits


def render_bundle(spec: dict, output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    mockup_dir = output_dir / "mockups"
    spec_path = output_dir / f"{sanitize_name(spec['safe_software_name'])}_spec.json"
    spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
    render_mockups(spec, mockup_dir)

    code_lines = build_code_lines(spec)
    source_line_count = sum(1 for line in code_lines if line.strip())

    file_stub = sanitize_name(spec["software_name"])
    application_path = output_dir / f"基于SpringBoot的{file_stub}_申请表.docx"
    manual_path = output_dir / f"基于Java&Vue的{file_stub}_操作手册.docx"
    code_path = output_dir / f"基于Java&Vue的{file_stub}_代码文档.docx"

    application_doc = Document(APPLICATION_TEMPLATE)
    set_header_text(application_doc, f"基于SpringBoot的{spec['software_name']} {spec['version']} 申请表")
    fill_application_tables(application_doc, spec, source_line_count)
    application_doc.save(application_path)

    manual_doc = Document(MANUAL_TEMPLATE)
    set_header_text(manual_doc, f"基于Java&Vue的{spec['software_name']} {spec['version']} 操作手册")
    rewrite_manual_paragraphs(manual_doc, build_manual_sequence(spec))
    manual_doc.save(manual_path)
    replace_docx_media(manual_path, mockup_dir)

    code_doc = Document(CODE_TEMPLATE)
    set_header_text(code_doc, f"基于Java&Vue的{spec['software_name']} {spec['version']} - 代码文档")
    write_code_document(code_doc, code_lines)
    code_doc.save(code_path)

    residue = {
        str(path): scan_for_residue(path)
        for path in (application_path, manual_path, code_path)
    }
    bad = {path: hits for path, hits in residue.items() if hits}
    if bad:
        raise RuntimeError(f"Template residue detected: {json.dumps(bad, ensure_ascii=False)}")

    return {
        "spec_path": str(spec_path),
        "application_path": str(application_path),
        "manual_path": str(manual_path),
        "code_path": str(code_path),
        "mockup_dir": str(mockup_dir),
        "source_line_count": source_line_count,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Render a software copyright draft bundle.")
    parser.add_argument("--name", help="Software full name")
    parser.add_argument("--intro", help="Short software introduction")
    parser.add_argument("--version", default="V1.0")
    parser.add_argument("--spec", help="Existing spec JSON path")
    parser.add_argument("--output-dir", required=True, help="Directory for rendered files")
    args = parser.parse_args()

    if args.spec:
        spec = json.loads(Path(args.spec).read_text(encoding="utf-8"))
    else:
        if not args.name or not args.intro:
            parser.error("--name and --intro are required when --spec is not provided.")
        spec = build_spec(args.name, args.intro, args.version)

    result = render_bundle(spec, Path(args.output_dir).resolve())
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
